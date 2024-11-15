def generate_surrogacy_agreement(name, text, agreeing_users):
    import uuid, os
    from django.conf import settings
    folder = 'surrogacy'
    HEIGHT = 11
    WIDTH = 8.5
    output_name = name + '-' + str(uuid.uuid4())
    base_dir = os.path.join(settings.BASE_DIR, 'media/{}/'.format(folder))
    code_lines = 45
    code_per_line = 90

    font_size = 13

    from PIL import Image
    from docx import Document
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Cm, Pt
    from pygments import highlight
    from pygments.lexers import PythonLexer, HtmlLexer, BashLexer, JavascriptLexer
    from pygments.formatters import ImageFormatter
    import re

    document = Document()

    text = ''

    title = text.split('\n')[0]

    document.add_heading(title, 0)

    text = text.replace('‘','\'').replace('’','\'')
    text_split = text.split('***')

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    add_page_number(document.sections[0].footer.paragraphs[0].add_run())

    sections = document.sections
    for section in sections:
        section.page_height = Inches(HEIGHT)
        section.page_width = Inches(WIDTH)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Cm(0.01)
    paragraph_format.space_after = Cm(0.01)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(font_size)
    signature_rep = "__________________________________"
    roles = ['Intended Parent', 'Intended Parent', 'Surrogate Mother', 'Surrogate Mother\'s Partner']
    signature_count = 0
    def add_paragraph(line):
        if signature_rep in line and agreeing_users[signature_count]:
            p = document.add_paragraph()
            run = p.add_run()
            run.add_text('X ')
            from jsignature.utils import draw_signature
            image_path = base_dir + 'signature-{}'.format(uuid.uuid4())
            file = draw_signature(agreeing_users[signature_count].verifications.last().signature)
            file.save(image_path)
            width = Image.open(image_path).size[0] / 90
                        if width > WIDTH - 1: width = WIDTH - 1
                        run.add_picture(image_path, width=Inches(width))
            signature_count = signature_count + 1
            run.add_text(', {} - Dated: {}'.format(roles[signature_count], timezone.now().strftime('%B, %d, %Y')))
        else:
            paragraph = document.add_paragraph(line)
            paragraph.style = document.styles['Normal']

    image_count = 0

    for t in text_split:
        split = re.split('\*[\w\.]+\*', t)
        language = '\n'
        try:
            language = t[len(split[0]):len(t)-len(split[1])][1:-1].lower()
        except: pass
        for line in split[0].split('\n'):
            paragraph = add_paragraph(line)
        code = split[1] if len(split) > 1 else False
        if code:
            run = True
            while run:
                s = code.split('\n')[:code_lines]
                lines_formatted = []
                for code_line in s:
                    for x in range(0, int(len(code_line)/code_per_line)):
                        lines_formatted = lines_formatted + [('(continued line) ' if x > 0 else '') + code_line[x*code_per_line:(x+1)*code_per_line]]
                c = '\n'.join(lines_formatted)
                remaining_code = '\n'.join(code.split('\n')[code_lines:])
                image_path = base_dir + 'image{}.png'.format(image_count)
                with open(image_path, "wb") as f:
                    add = True
                    print(language)
                    if language == 'python':
                        f.write(highlight(c, PythonLexer(), ImageFormatter()))
                    elif language == 'javascript':
                        f.write(highlight(c, JavascriptLexer(), ImageFormatter()))
                    elif language == 'html':
                        f.write(highlight(c, HtmlLexer(), ImageFormatter()))
                    elif language == 'bash':
                        f.write(highlight(c, BashLexer(), ImageFormatter()))
                    elif language.startswith('screenshot'):
                        image_path = base_dir + language
                    else:
                        add = False
                        for line in c.split('\n'):
                            paragraph = document.add_paragraph(line)
                            paragraph.style = document.styles['Normal']
                    f.close()
                    if add:
                        width = Image.open(image_path).size[0] / 90
                        if width > WIDTH - 1: width = WIDTH - 1
                        document.add_picture(image_path, width=Inches(width))
                    image_count = image_count + 1
                    if len(remaining_code) == 0: run = False
    savename = base_dir + '{}.docx'.format(output_name)
    document.save(savename)
    from spire.doc import *
    from spire.doc.common import *
    document = Document()
    document.LoadFromFile(savename)
    document.SaveToFile(savename + '.pdf', FileFormat.PDF)
    document.Close()
    return savename
