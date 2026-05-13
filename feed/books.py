def generate_post_book(post):
    from django.conf import settings
    import os, uuid
    path = os.path.join(settings.BASE_DIR, 'media/books/', '{}.docx'.format(str(uuid.uuid4())))
    return generate_book(post.content, path)


def generate_code_book(value, lang='en', src='en'):
    from django.conf import settings
    if not value: return value
    if not settings.USE_PRISM: return value
    import re, html
    from django.utils.html import strip_tags
    op = []
    title = value.split('\n')[0]
    v = re.split('(```)', value.replace('‘','\'').replace('’','\'')) #.split('```')
    language = ''
    nextislang = False
    nextiscode = False
    language = ''
    codeortext = False
    for t in v:
        if t == '': continue
        if t == '```':
            codeortext = not codeortext
            continue
        language = ''
        text = ''
        code = ''
        try:
            if codeortext:
                language = t.split('\n')[0].replace('\r', '') if len(t) > 0 else False
                code = t.split('\n', 1)[1] if len(t.split('\n', 1)[1]) > 0 else False
            else:
                text = t
        except: pass
        if language == 'html': language = 'markup'
        if code:
            lines = []
            for line in code.split('\n'):
                if len(line.rsplit('#', 1)) > 1:
                    to_trans = line.rsplit('#', 1)[1]
                    translated = translate(request, to_trans, target=lang, src=src)
                    line_string = line.rsplit('#', 1)[0] + '# ' + translated
                else: line_string = line
                lines = lines + [line_string]
            out = '\n'.join(lines)
            op = op + [{'text': translate(None, strip_tags(split[0]), target=lang, src=src), 'lang': language, 'code': html.escape(out) if language != 'markup' else '<!-- {} -->'.format(out)}]
        elif text:
            op = op + [{'text': translate(None, strip_tags(split[0]), target=lang, src=src)}]
    from django.template.loader import render_to_string
    return render_to_string('feed/book.html', {'value': op}), title

import textwrap

def split_text_parts(text, max_len=310, last_max_len=293):
    words = text.split()
    parts = []
    curr_part = []

    for word in words:
        # If adding the word exceeds limit for all except last
        tentative = ' '.join(curr_part + [word])
        # For all but last, use max_len; for last use last_max_len
        curr_max = last_max_len if len(parts) and len(' '.join(parts + [tentative])) <= last_max_len else max_len

        if len(tentative) > curr_max:
            # Finish current part and start a new one
            if curr_part:
                parts.append(' '.join(curr_part))
            curr_part = [word]
        else:
            curr_part.append(word)

    # Append any remaining words
    if curr_part:
        parts.append(' '.join(curr_part))

    # Now, if last part is too long, move words to previous
    while len(parts) > 1 and len(parts[-1]) >= last_max_len:
        # Move last word from penultimate to last
        last_words = parts[-1].split()
        prev_words = parts[-2].split()
        moved_word = prev_words.pop()
        parts[-2] = ' '.join(prev_words)
        parts[-1] = moved_word + ' ' + ' '.join(last_words)

    # Clean up any blanks
    parts = [part for part in parts if part.strip()]
    # Final check
    assert all(len(part) < max_len for part in parts[:-1]), "A non-last part exceeds 310 chars"

    return parts


def split_at_interval(arr, interval):
    return [arr[i : i + interval] for i in range(0, len(arr), interval)]

def generate_book(text, out_path_docx):
    from autocorrect import Speller
    speller = Speller()
    replace = {
        'Tango': 'Django',
        'request.OST': 'request.POST',
        'EMI_ADDRESS': 'EMAIL_ADDRESS',
        'EMI_HST_PASSWORD': 'EMAIL_HOST_PASSWORD',
        'SON': 'JSON',
        'Teilif.com': 'Twilio.com',
        'Teilif': 'Twilio',
    #    '': '',
    }
    def spell(line):
        text = speller(line)
        import re
        for key, value in replace.items():
            text = re.sub('\s' + key + '\s', ' ' + value + ' ', text)
            text = re.sub('\s' + key + '\.', ' ' + value + '.', text)
            text = re.sub('\s' + key + '\,', ' ' + value + ',', text)
        text = re.sub('(?P<get>\\.|!|;|:)[ \t]+(?P<put>\\w)', '\\g<get> \\g<put>', text)
        return text

    from docx import Document
#    from htmldocx import HtmlToDocx
    document = Document()
    import uuid, os
    from django.conf import settings
    from pygments import highlight
    from pygments.lexers import PythonLexer, HtmlLexer, BashLexer, JavascriptLexer
    from pygments.formatters import ImageFormatter
    HEIGHT = 9
    WIDTH = 6

    font_size = 13

    code_lines = 60
    code_per_line = 60

    base_dir = ''

    from PIL import Image
    from docx import Document
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Cm, Pt
    import re

    document = Document()
    title = text.split('\n')[0]

    document.add_heading(title, 0)

    text = text.replace('‘','\'').replace('’','\'')

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    add_page_number(document.sections[0].footer.paragraphs[0].add_run())

    sections = document.sections
    for section in sections:
        section.page_height = Inches(HEIGHT)
        section.page_width = Inches(WIDTH)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Cm(0.01)
    paragraph_format.space_after = Cm(0.01)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(font_size)

    max_width = Inches(WIDTH-(0.8*2))

    def add_paragraph(line):
        from django.utils.html import strip_tags
        import markdown
        paragraph = document.add_paragraph(strip_tags(markdown.markdown(line)))
        paragraph.style = document.styles['Normal']

    image_count = 0

    images = []
    import re
    v = re.split('(```)', text) #.split('```')
    language = ''
    codeortext = False
    for t in v:
        if t == '': continue
        if t == '```':
            codeortext = not codeortext
            continue
#        if t == '```' and not (nextislang or nextiscode):
#            nextislang = True
#            nextiscode = False
#            language = ''
#            text = ''
#            code = ''
#            continue
#        elif t == '```' and nextislang:
#            nextislang = False
#            continue
        language = ''
        text = ''
        code = ''
        try:
            if codeortext:
                language = t.split('\n')[0].replace('\r', '') if len(t) > 0 else False
                code = t.split('\n', 1)[1] if len(t.split('\n', 1)[1]) > 0 else False
            else:
                text = t
                for line in text.split('\n'):
                    paragraph = add_paragraph(spell(line))
        except: pass
        if code:
            run = True
            while run:
                s = code.split('\n')[:code_lines]
                code = '\n'.join(code.split('\n')[code_lines:])
                lines_formatted = []
                for code_line in s:
                    lines = textwrap.wrap(code_line, width=code_per_line, break_long_words=True)
                    first = True
                    count = 0
                    for line in lines:
                        if not first:
                            lines[count] == '→ ' + line
                        else:
                            first = False
                            lines[count] = line
                        count += 1
                    lines_formatted.extend(lines)
                pieces = split_at_interval(lines_formatted, code_lines)
                import math
                for x in range(0, len(pieces)):
                    c = '\n'.join(pieces[x])
                    image_path = base_dir + 'image-{}-{}.png'.format(str(uuid.uuid4()),image_count)
                    with open(image_path, "wb") as f:
                        add = True
                        print(language)
                        if language == 'python':
                            f.write(highlight(c, PythonLexer(), ImageFormatter()))
                            images = images + [image_path]
                        elif language == 'javascript':
                            f.write(highlight(c, JavascriptLexer(), ImageFormatter()))
                            images = images + [image_path]
                        elif language == 'html':
                            f.write(highlight(c, HtmlLexer(), ImageFormatter()))
                            images = images + [image_path]
                        elif language == 'bash':
                            f.write(highlight(c, BashLexer(), ImageFormatter()))
                            images = images + [image_path]
                        elif not language:
                            f.write(highlight(c, JavascriptLexer(), ImageFormatter()))
                            images = images + [image_path]
                        elif language.startswith('screenshot'):
                            image_path = base_dir + language
                            images = images + [image_path]
                        else:
                            add = False
                            for line in c.split('\n'):
                                paragraph = add_paragraph(spell(line))
                        f.close()
                        if add:
                            width = Image.open(image_path).size[0] / 120
                            height = Image.open(image_path).size[1] / 120
                            if height > HEIGHT - 3:
                                height = HEIGHT - 3
                                document.add_picture(image_path, height=Inches(height))
                            elif width > WIDTH - 3:
                                width = WIDTH - 3
                                document.add_picture(image_path, width=Inches(width))
                            else:
                                document.add_picture(image_path, width=Inches(width))
                        image_count = image_count + 1
                if len(code) == 0: run = False

    document.save(out_path_docx)
    for image in images:
        os.remove(image)
    return out_path_docx
